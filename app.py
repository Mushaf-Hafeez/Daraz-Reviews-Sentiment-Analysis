import streamlit as st
import joblib
import pandas as pd
import PyPDF2
from docx import Document
import tempfile
import os
from utils import preprocess_text

# Page configuration
st.set_page_config(
    page_title="Sentiment Analysis",
    page_icon="🔮",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Title and description
st.title("🔮 Binary Sentiment Analysis")
st.markdown("### Daraz E-Commerce Review Sentiment Classifier")
st.markdown("---")

# Sidebar information
with st.sidebar:
    st.header("ℹ️ About")
    st.markdown("""
    **Project**: Binary Sentiment Analysis
    
    **Model**: SVM (Support Vector Machine)
    
    **Features**:
    - TF-IDF Vectorization
    - 5000 features
    - Linear Kernel
    
    **Classes**:
    - 😞 Negative (-1)
    - 😊 Positive (1)
    
    **Language**: English + Roman Urdu (Code-mixed)
    """)
    
    st.markdown("---")
    st.markdown("**Features:**")
    st.markdown("""
    1. **Single Review Classification**
       - Enter one review
       - Get instant sentiment prediction
    
    2. **Batch File Upload**
       - Upload PDF, DOC, or DOCX
       - Analyze all reviews at once
       - Get percentage distribution
    
    **Supported Formats**: PDF, DOC, DOCX
    """)


# Load model and vectorizer
@st.cache_resource
def load_model():
    """Load the trained SVM model and vectorizer"""
    try:
        model = joblib.load('svm_model.pkl')
        vectorizer = joblib.load('vectorizer.pkl')
        return model, vectorizer
    except FileNotFoundError:
        st.error("❌ Model files not found! Please run the main.ipynb notebook first to train and save the model.")
        st.stop()

# Load model
svm_model, vectorizer = load_model()

# Helper functions for file processing
def extract_text_from_pdf(pdf_file):
    """Extract text from PDF file"""
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(docx_file):
    """Extract text from DOCX file"""
    doc = Document(docx_file)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
    return text

def extract_text_from_doc(doc_file):
    """Extract text from DOC file using python-docx (basic support)"""
    # For .doc files, we'll try to use python-docx if possible
    try:
        doc = Document(doc_file)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except:
        st.warning("Note: Basic .doc support. For best results, please use .docx files.")
        return ""

def parse_reviews_from_text(text):
    """Parse reviews from extracted text (one review per line)"""
    reviews = [line.strip() for line in text.split('\n') if line.strip() and len(line.strip()) > 10]
    return reviews

def classify_reviews_batch(reviews):
    """Classify multiple reviews and return statistics"""
    predictions = []
    
    for review in reviews:
        try:
            processed = preprocess_text(review)
            vectorized = vectorizer.transform([processed])
            prediction = svm_model.predict(vectorized)[0]
            predictions.append(prediction)
        except:
            continue
    
    return predictions

# Create tabs for different functionalities
tab1, tab2 = st.tabs(["📝 Single Review", "📤 Batch Upload"])

# TAB 1: Single Review Classification
with tab1:
    st.markdown("### 📝 Classify Single Review")
    
    # Create input box
    review_input = st.text_area(
        "Paste your product review here:",
        placeholder="Example: This product is amazing! It works perfectly and arrived quickly.",
        height=150,
        label_visibility="collapsed"
    )
    
    # Create columns for buttons
    col1, col2 = st.columns([3, 1])
    
    with col2:
        predict_button = st.button("Classify", type="primary", use_container_width=True)
    
    # Process and predict
    if predict_button:
        if not review_input.strip():
            st.warning("⚠️ Please enter a review before clicking Classify")
        else:
            with st.spinner("🔄 Processing your review..."):
                # Preprocess the text
                processed_review = preprocess_text(review_input)
                
                # Vectorize
                vectorized_input = vectorizer.transform([processed_review])
                
                # Make prediction
                prediction = svm_model.predict(vectorized_input)[0]
                
                # Get confidence scores (decision function)
                confidence_scores = svm_model.decision_function(vectorized_input)[0]
                
                # Map prediction to sentiment (BINARY ONLY)
                sentiment_map = {-1: "Negative", 1: "Positive"}
                emoji_map = {-1: "😞", 1: "😊"}
                
                sentiment_label = sentiment_map[prediction]
                emoji = emoji_map[prediction]
                
                # Display results
                st.markdown("---")
                st.markdown("### 🎯 Prediction Results")
                
                # Sentiment display
                if prediction == 1:
                    st.success(f"**Sentiment:** {emoji} **{sentiment_label.upper()}** (Class: {prediction})")
                else:
                    st.error(f"**Sentiment:** {emoji} **{sentiment_label.upper()}** (Class: {prediction})")
                
                # Processed review
                with st.expander("📖 View Processed Review"):
                    st.write(processed_review)
                
                # Model information
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Model", "SVM")
                with col2:
                    st.metric("Vectorizer", "TF-IDF")
                with col3:
                    st.metric("Features", "5000")
                
                st.markdown("---")
                st.success("✅ Classification complete!")

# TAB 2: Batch File Upload
with tab2:
    st.markdown("### 📤 Upload File for Batch Analysis")
    st.markdown("Upload a PDF, DOC, or DOCX file containing reviews (one review per line)")
    
    # File uploader
    uploaded_file = st.file_uploader(
        "Choose a file",
        type=["pdf", "doc", "docx"],
        help="Upload PDF, DOC, or DOCX files containing reviews"
    )
    
    if uploaded_file is not None:
        st.markdown("---")
        
        # Extract text based on file type
        with st.spinner("📖 Extracting text from file..."):
            file_extension = uploaded_file.name.split(".")[-1].lower()
            
            try:
                if file_extension == "pdf":
                    extracted_text = extract_text_from_pdf(uploaded_file)
                elif file_extension == "docx":
                    extracted_text = extract_text_from_docx(uploaded_file)
                elif file_extension == "doc":
                    extracted_text = extract_text_from_doc(uploaded_file)
                else:
                    st.error("Unsupported file format!")
                    extracted_text = ""
                
                if extracted_text:
                    # Parse reviews
                    reviews = parse_reviews_from_text(extracted_text)
                    
                    st.info(f"📊 Found **{len(reviews)}** reviews in the file")
                    
                    if len(reviews) == 0:
                        st.warning("⚠️ No reviews found in the file. Please ensure reviews are separated by line breaks and contain more than 10 characters.")
                    else:
                        # Classify button
                        if st.button("🚀 Analyze All Reviews", type="primary", use_container_width=True):
                            with st.spinner(f"🔄 Classifying {len(reviews)} reviews..."):
                                # Classify all reviews
                                predictions = classify_reviews_batch(reviews)
                                
                                if len(predictions) > 0:
                                    # Calculate statistics
                                    negative_count = predictions.count(-1)
                                    positive_count = predictions.count(1)
                                    total = len(predictions)
                                    
                                    negative_percentage = (negative_count / total) * 100
                                    positive_percentage = (positive_count / total) * 100
                                    
                                    # Display statistics
                                    st.markdown("---")
                                    st.markdown("### 📊 Analysis Results")
                                    
                                    # Main metrics
                                    col1, col2, col3 = st.columns(3)
                                    
                                    with col1:
                                        st.metric("Total Reviews Analyzed", total)
                                    with col2:
                                        st.metric("😞 Negative Reviews", f"{negative_count} ({negative_percentage:.1f}%)")
                                    with col3:
                                        st.metric("😊 Positive Reviews", f"{positive_count} ({positive_percentage:.1f}%)")
                                    
                                    # Sentiment distribution chart
                                    st.markdown("---")
                                    st.markdown("### 📈 Sentiment Distribution")
                                    
                                    # Create pie chart
                                    sentiment_data = {
                                        '😞 Negative': negative_count,
                                        '😊 Positive': positive_count
                                    }
                                    
                                    col1, col2 = st.columns([1, 2])
                                    
                                    with col1:
                                        st.markdown("**Summary:**")
                                        st.markdown(f"""
                                        - **Negative**: {negative_percentage:.1f}%
                                        - **Positive**: {positive_percentage:.1f}%
                                        """)
                                    
                                    with col2:
                                        # Create a simple bar chart
                                        chart_data = pd.DataFrame({
                                            'Sentiment': ['😞 Negative', '😊 Positive'],
                                            'Count': [negative_count, positive_count],
                                            'Percentage': [negative_percentage, positive_percentage]
                                        })
                                        st.bar_chart(chart_data.set_index('Sentiment')['Count'])
                                    
                                    # Detailed results table
                                    st.markdown("---")
                                    st.markdown("### 📋 Detailed Results")
                                    
                                    results_df = pd.DataFrame({
                                        'Review': reviews[:10],
                                        'Prediction': [sentiment_map.get(p, "Unknown") for p in predictions[:10]]
                                    })
                                    
                                    st.dataframe(results_df, use_container_width=True)
                                    
                                    if len(reviews) > 10:
                                        st.info(f"Showing first 10 reviews out of {len(reviews)}")
                                    
                                    st.markdown("---")
                                    st.success("✅ Analysis complete!")
                                    
                                    # Download results button
                                    results_df_full = pd.DataFrame({
                                        'Review': reviews,
                                        'Prediction': [sentiment_map.get(p, "Unknown") for p in predictions]
                                    })
                                    
                                    csv = results_df_full.to_csv(index=False)
                                    st.download_button(
                                        label="📥 Download Results as CSV",
                                        data=csv,
                                        file_name="sentiment_analysis_results.csv",
                                        mime="text/csv"
                                    )
                                else:
                                    st.error("❌ Error classifying reviews. Please try again.")
                else:
                    st.error("❌ Could not extract text from file. Please ensure the file is valid.")
                    
            except Exception as e:
                st.error(f"❌ Error processing file: {str(e)}")

# Footer
st.markdown("---")
st.markdown("**Binary Sentiment Analysis System** | Negative (-1) & Positive (1) Classification")
